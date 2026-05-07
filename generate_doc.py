import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = docx.Document()

# Add a Title
title = doc.add_heading('CÁC TÍNH NĂNG NỔI BẬT TRONG DỰ ÁN GALAXY STUDIO', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph('Dưới đây là mô tả chi tiết về 3 tính năng nổi bật được xây dựng sát với thực tế doanh nghiệp trong dự án quản lý rạp chiếu phim (Galaxy Studio):\n')

# Feature 1
h1 = doc.add_heading('1. Xử lý đồng thời chống đặt trùng ghế (Concurrency Control & Transaction)', level=1)
p1_1 = doc.add_paragraph()
p1_1.add_run('Mô tả kỹ thuật: ').bold = True
p1_1.add_run('Đây là tính năng giải quyết bài toán "Race Condition" kinh điển khi nhiều khách hàng truy cập và chọn cùng một vị trí ghế ngồi trong cùng một suất chiếu tại cùng một thời điểm.')

p1_2 = doc.add_paragraph()
p1_2.add_run('Cách thức hoạt động: ').bold = True
p1_2.add_run('Khi người dùng bắt đầu thao tác thanh toán, hệ thống Backend (Node.js/Express) sẽ khởi tạo một Database Transaction (Giao dịch cơ sở dữ liệu) thông qua MySQL2. Toàn bộ các thao tác như tạo hóa đơn, tạo vé, cập nhật trạng thái thanh toán và cộng/trừ điểm tích lũy đều được đưa vào cùng một Transaction duy nhất. Nếu có bất kỳ sự cố nào xảy ra (ví dụ có người khác đã nhanh tay thanh toán ghế đó), toàn bộ chuỗi thao tác sẽ bị hủy bỏ (Rollback) và không ghi nhận vào Database.')

p1_3 = doc.add_paragraph()
p1_3.add_run('Giá trị: ').bold = True
p1_3.add_run('Đảm bảo tính toàn vẹn dữ liệu (ACID) của hệ thống, tránh rủi ro bán một ghế cho nhiều người gây xung đột trong vận hành rạp, thể hiện một kiến trúc backend vững chắc.')

# Feature 2
h2 = doc.add_heading('2. Chấm công bằng ảnh chụp Camera trực tiếp tại rạp (HTML5 Webcam API)', level=1)
p2_1 = doc.add_paragraph()
p2_1.add_run('Mô tả kỹ thuật: ').bold = True
p2_1.add_run('Một giải pháp chống gian lận chấm công (chấm công hộ) hiện đại, thay thế cho việc quẹt thẻ từ hoặc vân tay truyền thống bằng cách sử dụng chính Camera của thiết bị web.')

p2_2 = doc.add_paragraph()
p2_2.add_run('Cách thức hoạt động: ').bold = True
p2_2.add_run('Ở phía Frontend (ReactJS), hệ thống sử dụng Web API chuẩn (navigator.mediaDevices.getUserMedia) để kích hoạt Webcam hoặc Camera điện thoại. Hình ảnh luồng video (stream) được truyền lên giao diện, người dùng bấm nút chụp, hệ thống sẽ sử dụng thẻ <canvas> để bắt lại khung hình và mã hóa thành chuỗi Base64. Chuỗi ảnh này lập tức được gửi tới Backend thông qua API để lưu trực tiếp vào cơ sở dữ liệu cùng với thời gian Check-in/Check-out.')

p2_3 = doc.add_paragraph()
p2_3.add_run('Giá trị: ').bold = True
p2_3.add_run('Giúp quản trị viên dễ dàng rà soát tính trung thực của nhân viên dựa vào ảnh chụp hiện trường tại thời điểm chấm công. Thể hiện khả năng tương tác sâu với phần cứng thiết bị từ ứng dụng Web.')

# Feature 3
h3 = doc.add_heading('3. Tự động gửi Email vé điện tử (E-Ticket) đính kèm mã QR', level=1)
p3_1 = doc.add_paragraph()
p3_1.add_run('Mô tả kỹ thuật: ').bold = True
p3_1.add_run('Tính năng tự động hóa luồng nghiệp vụ cung cấp vé điện tử, tích hợp hệ thống thư điện tử và mã vạch bảo mật, giúp giảm thiểu việc sử dụng vé giấy truyền thống.')

p3_2 = doc.add_paragraph()
p3_2.add_run('Cách thức hoạt động: ').bold = True
p3_2.add_run('Ngay sau khi Backend lưu trữ giao dịch Transaction đặt vé thành công, hệ thống sử dụng thư viện tạo mã "qrcode" để sinh ra một mã QR duy nhất chứa thông tin mã hóa của vé. Tiếp đó, module "nodemailer" được kích hoạt kết nối với SMTP Server. Backend sẽ tự động biên soạn một email bằng mã HTML chuyên nghiệp với đầy đủ thông tin rạp, giờ chiếu, số ghế, combo bắp nước và đính kèm trực tiếp mã QR này gửi thẳng vào hộp thư của người mua vé.')

p3_3 = doc.add_paragraph()
p3_3.add_run('Giá trị: ').bold = True
p3_3.add_run('Khách hàng chỉ cần mang mã QR trên điện thoại đến rạp để nhân viên soát vé quét (Scan) đi thẳng vào phòng chiếu. Điều này tạo ra một luồng trải nghiệm khách hàng (User Experience) hoàn toàn liền mạch và số hóa.')

# Save Document
doc.save('Cac_Tinh_Nang_Noi_Bat_GalaxyStudio.docx')
print("Document updated successfully.")
